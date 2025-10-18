import docx
import json
import os
import re

# Read the document
doc = docx.Document("AI Use case catalog.docx")

# Extract all text with formatting
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append({
            'text': para.text,
            'style': para.style.name
        })

# Also check for tables
tables_data = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        table_data.append(row_data)
    tables_data.append(table_data)

# Save to JSON for analysis
output = {
    'paragraphs': full_text,
    'tables': tables_data
}

with open('docx_content.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, indent=2, ensure_ascii=False)

print(f"Extracted {len(full_text)} paragraphs and {len(tables_data)} tables")
print("\nFirst 20 paragraphs:")
for i, para in enumerate(full_text[:20]):
    print(f"{i}: [{para['style']}] {para['text'][:100]}")
